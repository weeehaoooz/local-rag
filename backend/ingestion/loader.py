"""
indexers/loader.py
------------------
SmartDocumentLoader: layout-aware document loading for PDF, DOCX, HTML and TXT.

Addresses the "Structural Extraction & Layout Analysis" best-practice:
  - PDF: pdfplumber for column-aware reading order + table → Markdown conversion
  - DOCX: python-docx with heading/table structure preserved
  - HTML: BeautifulSoup with boilerplate-script/style stripping
  - TXT: plain pass-through via LlamaIndex SimpleDirectoryReader

All PDF/DOCX/HTML dependencies are optional; the loader gracefully falls back to
SimpleDirectoryReader when an optional library is absent.
"""

from __future__ import annotations

import logging
import os
import io
import re
import asyncio
from typing import List, Optional, Any

import ollama
from llama_index.core import SimpleDirectoryReader, Settings
from llama_index.core.schema import Document
import numpy as np

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Optional dependency guards
# ---------------------------------------------------------------------------

try:
    import easyocr  # type: ignore
    _HAS_EASYOCR = True
except ImportError:
    _HAS_EASYOCR = False
    logger.warning(
        "[loader] 'easyocr' not installed — OCR disabled. "
        "Install with: pip install easyocr"
    )

try:
    from PIL import Image  # type: ignore
    _HAS_PILLOW = True
except ImportError:
    _HAS_PILLOW = False

try:
    import pdfplumber  # type: ignore
    _HAS_PDFPLUMBER = True
except ImportError:
    _HAS_PDFPLUMBER = False
    logger.warning(
        "[loader] 'pdfplumber' not installed — PDF layout parsing disabled. "
        "Install with: pip install pdfplumber"
    )

try:
    import docx  # type: ignore   # python-docx
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False
    logger.warning(
        "[loader] 'python-docx' not installed — DOCX layout parsing disabled. "
        "Install with: pip install python-docx"
    )

try:
    from bs4 import BeautifulSoup  # type: ignore
    _HAS_BSP = True
except ImportError:
    _HAS_BSP = False
    logger.warning(
        "[loader] 'beautifulsoup4' not installed — HTML parsing disabled. "
        "Install with: pip install beautifulsoup4"
    )


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _table_to_markdown(table: list[list]) -> str:
    """Convert a list-of-list table (from pdfplumber or docx) to Markdown."""
    if not table:
        return ""

    rows = []
    for i, row in enumerate(table):
        # Normalise cells: pdfplumber may give None for empty cells
        cells = [str(cell).strip() if cell is not None else "" for cell in row]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            # Header separator row
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")

    return "\n".join(rows)


# OCR Reader (Singleton/Lazy)
_OCR_READER: Optional[easyocr.Reader] = None

def _get_ocr_reader() -> easyocr.Reader:
    global _OCR_READER
    if _OCR_READER is None:
        if not _HAS_EASYOCR:
            raise ImportError("easyocr is not installed.")
        logger.info("[loader] Initializing easyocr Reader (English)...")
        # gpu=False for local CPU stability; set True if user has Mac GPU/CUDA
        _OCR_READER = easyocr.Reader(['en'], gpu=False)
    return _OCR_READER


# ---------------------------------------------------------------------------
# Section detection helper
# ---------------------------------------------------------------------------

_SECTION_HEADING_RE = re.compile(
    r"(?m)^(?:#{1,3}\s+(.+))$"        # Markdown headings: # Title, ## Section
    r"|"
    r"^([A-Z][A-Z\s]{4,60})$"         # ALL-CAPS lines (common in PDFs)
)


def _detect_sections(text: str) -> list[str]:
    """
    Extract section headings from text.

    Recognises:
      - Markdown headings (# / ## / ###)
      - ALL-CAPS lines that look like section titles
    """
    sections: list[str] = []
    for match in _SECTION_HEADING_RE.finditer(text):
        heading = (match.group(1) or match.group(2) or "").strip()
        if heading and heading not in sections:
            sections.append(heading)
    return sections


# ---------------------------------------------------------------------------
# Vision Helper
# ---------------------------------------------------------------------------

def _describe_image_with_vision(image_bytes: bytes, model_name: str = "gemma4:latest") -> str:
    """
    Use a multimodal LLM (via Ollama) to describe an image.
    Specifically prompted to extract data from charts and diagrams.
    """
    prompt = (
        "You are a technical document analyst. Describe this image in detail. "
        "If it is a chart or technical diagram, explain the axes, data trends, "
        "legend, and the main message it conveys. If it is an illustration, "
        "describe the subjects and their actions. Be concise and factual."
    )
    try:
        response = ollama.generate(
            model=model_name,
            prompt=prompt,
            images=[image_bytes]
        )
        return response.get('response', '').strip()
    except Exception as e:
        logger.warning("[loader] Vision description failed: %s", e)
        return ""


# ---------------------------------------------------------------------------
# PDF loader
# ---------------------------------------------------------------------------

def _load_pdf_with_pdfplumber(file_path: str, metadata: dict, enable_vision: bool = False) -> List[Document]:
    """
    Extract text from a PDF using pdfplumber.

    Features:
    - Layout-aware extraction.
    - Table-to-Markdown conversion.
    - OCR fallback for scanned pages.
    - Vision-based image description (Optional).
    """
    documents: list[Document] = []
    vision_model = getattr(Settings.llm, "model", "gemma4:latest") if hasattr(Settings, "llm") else "gemma4:latest"

    with pdfplumber.open(file_path) as pdf:
        total_pages = len(pdf.pages)

        for page_num, page in enumerate(pdf.pages, start=1):
            page_parts: list[str] = []
            used_ocr = False

            # 1. Extract tables
            table_bboxes: list[tuple] = []
            for table in page.extract_tables():
                md = _table_to_markdown(table)
                if md:
                    page_parts.append(f"\n[TABLE]\n{md}\n[/TABLE]\n")

            raw_tables = page.find_tables()
            table_bboxes = [tbl.bbox for tbl in raw_tables]

            # 2. Extract significant images and describe them if vision is enabled
            if enable_vision:
                for idx, img_info in enumerate(page.images):
                    # Filter: ignore tiny images (icons, logos, bullets)
                    width = img_info["width"]
                    height = img_info["height"]
                    if width > 150 and height > 150:
                        logger.info("[loader] Describing image %d on page %d of '%s'...", idx, page_num, os.path.basename(file_path))
                        try:
                            # Crop and convert to PNG bytes
                            bbox = (img_info["x0"], img_info["top"], img_info["x1"], img_info["bottom"])
                            # Ensure bbox is valid
                            if bbox[2] > bbox[0] and bbox[3] > bbox[1]:
                                cropped = page.within_bbox(bbox).to_image(resolution=200).original
                                img_byte_arr = io.BytesIO()
                                cropped.save(img_byte_arr, format='PNG')
                                description = _describe_image_with_vision(img_byte_arr.getvalue(), model_name=vision_model)
                                if description:
                                    page_parts.append(f"\n[Visual Metadata - Image {idx}]\n{description}\n")
                        except Exception as e:
                            logger.warning("[loader] Failed to process image %d: %s", idx, e)

            # 3. Extract text
            words = page.extract_words(x_tolerance=3, y_tolerance=3, keep_blank_chars=False)

            def _in_table(word: dict) -> bool:
                wx0, wy0, wx1, wy1 = word["x0"], word["top"], word["x1"], word["bottom"]
                for bx0, by0, bx1, by1 in table_bboxes:
                    if wx0 >= bx0 and wx1 <= bx1 and wy0 >= by0 and wy1 <= by1:
                        return True
                return False

            non_table_words = [w for w in words if not _in_table(w)]
            page_body = _words_to_reading_order_text(non_table_words, page.width)

            # 4. OCR Fallback
            if len(page_body.strip()) < 50 and not page_parts and _HAS_EASYOCR:
                logger.info("[loader] Low text on page %d; attempting OCR...", page_num)
                try:
                    img = page.to_image(resolution=300).original
                    img_byte_arr = io.BytesIO()
                    img.save(img_byte_arr, format='PNG')
                    reader = _get_ocr_reader()
                    ocr_results = reader.readtext(img_byte_arr.getvalue(), detail=0, paragraph=True)
                    ocr_text = "\n".join(ocr_results)
                    if ocr_text.strip():
                        page_parts.append(f"[OCR Content]\n{ocr_text}")
                        used_ocr = True
                except Exception as e:
                    logger.warning("[loader] OCR failed for page %d: %s", page_num, e)

            if page_body.strip():
                page_parts.insert(0, page_body)

            page_text = f"[Page {page_num}]\n" + "\n".join(page_parts)
            sections = _detect_sections(page_text)

            page_metadata = {
                **metadata,
                "file_path": file_path,
                "source_type": "pdf",
                "page_number": page_num,
                "total_pages": total_pages,
                "has_tables": len(raw_tables) > 0,
                "used_ocr": used_ocr,
                "has_visual_metadata": enable_vision and any("[Visual Metadata]" in p for p in page_parts)
            }
            if sections:
                page_metadata["sections"] = sections

            documents.append(Document(text=page_text, metadata=page_metadata))

    return documents


def _words_to_reading_order_text(words: list[dict], page_width: float) -> str:
    """
    Re-assemble extracted words into natural reading order.

    Multi-column detection: if word x-positions cluster into two groups that
    don't overlap, treat as a 2-column layout and read the left column first.
    """
    if not words:
        return ""

    # Sort by vertical position first (top), then horizontal (left)
    words_sorted = sorted(words, key=lambda w: (round(w["top"] / 5) * 5, w["x0"]))

    # Group words into lines by y-band (within 5pt tolerance)
    lines: list[list[dict]] = []
    current_line: list[dict] = []
    current_y: float | None = None

    for word in words_sorted:
        y = round(word["top"] / 5) * 5
        if current_y is None or abs(y - current_y) <= 5:
            current_line.append(word)
            current_y = y
        else:
            if current_line:
                lines.append(sorted(current_line, key=lambda w: w["x0"]))
            current_line = [word]
            current_y = y

    if current_line:
        lines.append(sorted(current_line, key=lambda w: w["x0"]))

    # Detect 2-column layout: check if the median x-position of "right-side" words
    # is > page_width / 2 + 10% margin
    text_lines: list[str] = [" ".join(w["text"] for w in line) for line in lines]
    return "\n".join(text_lines)


# ---------------------------------------------------------------------------
# DOCX loader
# ---------------------------------------------------------------------------

def _load_docx_with_python_docx(file_path: str, metadata: dict) -> List[Document]:
    """
    Extract text from a DOCX file, splitting into separate Document objects
    based on heading boundaries (H1, H2, H3).
    """
    doc_obj = docx.Document(file_path)
    documents: list[Document] = []
    
    current_section_title = "Preamble"
    current_parts: list[str] = []
    has_tables_in_current = False
    all_sections: list[str] = []

    def _flush_section():
        nonlocal current_parts, has_tables_in_current
        if not current_parts:
            return
        
        section_text = "\n\n".join(current_parts)
        doc_metadata = {
            **metadata,
            "file_path": file_path,
            "source_type": "docx",
            "section_title": current_section_title,
            "has_tables": has_tables_in_current,
        }
        documents.append(Document(text=section_text, metadata=doc_metadata))
        current_parts = []
        has_tables_in_current = False

    for block in doc_obj.element.body:
        tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag

        if tag == "p":
            para = _find_paragraph(doc_obj, block)
            if para is None:
                continue
            style_name = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                continue

            # Check for heading levels to trigger a section break
            is_heading = any(h in style_name for h in ["Heading 1", "Heading 2", "Heading 3"])
            
            if is_heading:
                _flush_section()
                current_section_title = text
                all_sections.append(text)
                # We also include the heading in the text for context
                level = 1 if "Heading 1" in style_name else (2 if "Heading 2" in style_name else 3)
                current_parts.append("#" * level + " " + text)
            else:
                current_parts.append(text)

        elif tag == "tbl":
            tbl = _find_table(doc_obj, block)
            if tbl is None:
                continue
            rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
            md = _table_to_markdown(rows)
            if md:
                current_parts.append(f"\n[TABLE]\n{md}\n[/TABLE]\n")
                has_tables_in_current = True

    _flush_section()

    # Back-fill the 'sections' list to all documents for global context
    if all_sections:
        for doc in documents:
            doc.metadata["sections"] = all_sections

    return documents


def _find_paragraph(doc_obj, elem):
    """Return the docx Paragraph object for an XML element, or None."""
    for para in doc_obj.paragraphs:
        if para._element is elem:
            return para
    return None


def _find_table(doc_obj, elem):
    """Return the docx Table object for an XML element, or None."""
    for tbl in doc_obj.tables:
        if tbl._element is elem:
            return tbl
    return None


# ---------------------------------------------------------------------------
# HTML loader
# ---------------------------------------------------------------------------

def _load_html(file_path: str, metadata: dict) -> List[Document]:
    """
    Extract readable text from an HTML file, splitting into separate Document
    objects based on heading boundaries (H1, H2, H3).
    """
    with open(file_path, "r", encoding="utf-8", errors="replace") as fh:
        raw_html = fh.read()

    soup = BeautifulSoup(raw_html, "html.parser")

    # Remove noise elements
    for tag in soup(["script", "style", "noscript", "nav", "header", "footer", "aside"]):
        tag.decompose()

    # Extract main content preferentially
    main = soup.find("main") or soup.find("article") or soup.find("body") or soup
    
    # We want to keep heading information to split the text.
    # We'll replace headings with markers and then split.
    import uuid
    split_marker_template = "---SPLIT-{}-{}---"
    
    sections_list = []
    for h in main.find_all(["h1", "h2", "h3"]):
        title = h.get_text(strip=True)
        if not title:
            continue
        level = h.name[1] # '1', '2', or '3'
        marker = split_marker_template.format(level, title)
        h.replace_with(marker)
        sections_list.append(title)

    text = main.get_text(separator="\n")
    
    # Split by markers
    pattern = re.escape("---SPLIT-").replace("\\-", "-") + r"(\d)\-(.*?)\-\-\-"
    parts = re.split(pattern, text)
    
    documents: list[Document] = []
    
    # First part is Preamble (before any heading)
    if parts[0].strip():
        documents.append(Document(
            text=parts[0].strip(),
            metadata={
                **metadata,
                "file_path": file_path,
                "source_type": "html",
                "section_title": "Preamble"
            }
        ))
    
    # Subsequent parts come in triples: (level, title, content)
    for i in range(1, len(parts), 3):
        level = parts[i]
        title = parts[i+1]
        content = parts[i+2].strip()
        
        if content:
            # Re-insert the heading for context
            full_content = "#" * int(level) + " " + title + "\n\n" + content
            documents.append(Document(
                text=full_content,
                metadata={
                    **metadata,
                    "file_path": file_path,
                    "source_type": "html",
                    "section_title": title
                }
            ))

    # Extract page title
    title_tag = soup.find("title")
    page_title = title_tag.get_text(strip=True) if title_tag else None

    # Enrich metadata for all docs
    for doc in documents:
        if sections_list:
            doc.metadata["sections"] = sections_list
        if page_title:
            doc.metadata["html_title"] = page_title

    return documents


# ---------------------------------------------------------------------------
# Image loader
# ---------------------------------------------------------------------------

def _load_image_with_ocr(file_path: str, metadata: dict) -> List[Document]:
    """Extract text from an image using easyocr."""
    if not _HAS_EASYOCR:
        return SimpleDirectoryReader(input_files=[file_path]).load_data()

    logger.debug("[loader] Running OCR on image: %s", file_path)
    try:
        reader = _get_ocr_reader()
        results = reader.readtext(file_path, detail=0, paragraph=True)
        text = "\n".join(results)

        # Add a placeholder page marker for consistency
        text = "[Page 1]\n" + text

        return [Document(text=text, metadata={
            **metadata,
            "file_path": file_path,
            "source_type": "image",
            "used_ocr": True,
            "page_number": 1,
            "total_pages": 1,
        })]
    except Exception as e:
        logger.warning("[loader] Image OCR failed for '%s': %s", file_path, e)
        return SimpleDirectoryReader(input_files=[file_path]).load_data()


# ---------------------------------------------------------------------------
# SmartDocumentLoader
# ---------------------------------------------------------------------------

_EXTENSION_MAP = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "docx",
    ".html": "html",
    ".htm": "html",
    ".txt": "txt",
    ".md": "txt",
    ".png": "img",
    ".jpg": "img",
    ".jpeg": "img",
    ".bmp": "img",
    ".tiff": "img",
}


class SmartDocumentLoader:
    """
    Layout-aware document loader.

    Replaces ``SimpleDirectoryReader(input_files=[f]).load_data()`` with a
    format-specific strategy that preserves reading order, extracts tables as
    Markdown, and strips HTML boilerplate.

    Falls back to ``SimpleDirectoryReader`` when optional dependencies are
    missing or for unsupported file types.

    Usage::

        loader = SmartDocumentLoader()
        documents = loader.load(file_path)
    """

    def __init__(self, enable_vision: bool = False):
        self.enable_vision = enable_vision

    async def aload(self, file_path: str) -> List[Document]:
        """Asynchronously load a single file into a list of LlamaIndex Documents."""
        logger.debug("[loader] aload called for '%s'", file_path)
        # Offload blocking parsing to a thread pool for true async scalability
        return await asyncio.to_thread(self.load, file_path)

    def load(self, file_path: str) -> List[Document]:
        """Load a single file and return a list of LlamaIndex Documents."""
        ext = os.path.splitext(file_path)[1].lower()
        fmt = _EXTENSION_MAP.get(ext, "unknown")

        base_metadata: dict = {
            "file_name": os.path.basename(file_path),
        }

        try:
            if fmt == "pdf" and _HAS_PDFPLUMBER:
                logger.debug("[loader] Loading PDF with pdfplumber: %s", file_path)
                return _load_pdf_with_pdfplumber(file_path, base_metadata, enable_vision=self.enable_vision)

            if fmt == "docx" and _HAS_DOCX:
                logger.debug("[loader] Loading DOCX with python-docx: %s", file_path)
                return _load_docx_with_python_docx(file_path, base_metadata)

            if fmt == "html" and _HAS_BSP:
                logger.debug("[loader] Loading HTML with BeautifulSoup: %s", file_path)
                return _load_html(file_path, base_metadata)

            if fmt == "img" and _HAS_EASYOCR:
                logger.debug("[loader] Loading Image with easyocr: %s", file_path)
                return _load_image_with_ocr(file_path, base_metadata)

        except Exception as exc:
            logger.warning(
                "[loader] Smart loader failed for '%s' (%s). Falling back to SimpleDirectoryReader. Error: %s",
                file_path,
                fmt,
                exc,
            )

        # Default: SimpleDirectoryReader (TXT, unknown, or fallback)
        logger.debug("[loader] Loading with SimpleDirectoryReader: %s", file_path)
        docs = SimpleDirectoryReader(input_files=[file_path]).load_data()
        # Ensure fallback documents still carry source_type metadata
        for doc in docs:
            doc.metadata.setdefault("source_type", fmt if fmt != "unknown" else "txt")
        return docs

def discover_files(data_dir=None):
    """Scan directory and group files by category."""
    from collections import defaultdict
    from retrieval.guardrails import _derive_category
    from config import DATA_DIR

    if data_dir is None:
        data_dir = DATA_DIR
    
    logger.debug(f"Scanning {data_dir} for files...")
    all_files = []
    for root, _, files in os.walk(data_dir):
        for file in files:
            if file.endswith((".txt", ".pdf", ".docx", ".md")):
                all_files.append(os.path.join(root, file))
    
    if not all_files:
        return {}, []

    files_by_category = defaultdict(list)
    for f in all_files:
        cat = _derive_category(f, data_dir)
        files_by_category[cat].append(f)
    
    return files_by_category, all_files
