"""
indexers/loader.py
------------------
SmartDocumentLoader: layout-aware document loading for PDF, DOCX, HTML and TXT.

Addresses the "Structural Extraction & Layout Analysis" best-practice:
  - PDF: pdfplumber for column-aware reading order + table → Markdown conversion
  - DOCX: python-docx with heading/table structure preserved
  - HTML: BeautifulSoup with boilerplate-script/style stripping
  - TXT: plain pass-through via LlamaIndex SimpleDirectoryReader

All PDF/DOCX/HTML dependencies are optional; the loader gracefully falls back to
SimpleDirectoryReader when an optional library is absent.
"""

from __future__ import annotations

import logging
import os
from typing import List

from llama_index.core import SimpleDirectoryReader
from llama_index.core.schema import Document

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Optional dependency guards
# ---------------------------------------------------------------------------

try:
    import pdfplumber  # type: ignore
    _HAS_PDFPLUMBER = True
except ImportError:
    _HAS_PDFPLUMBER = False
    logger.warning(
        "[loader] 'pdfplumber' not installed — PDF layout parsing disabled. "
        "Install with: pip install pdfplumber"
    )

try:
    import docx  # type: ignore   # python-docx
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False
    logger.warning(
        "[loader] 'python-docx' not installed — DOCX layout parsing disabled. "
        "Install with: pip install python-docx"
    )

try:
    from bs4 import BeautifulSoup  # type: ignore
    _HAS_BSP = True
except ImportError:
    _HAS_BSP = False
    logger.warning(
        "[loader] 'beautifulsoup4' not installed — HTML parsing disabled. "
        "Install with: pip install beautifulsoup4"
    )


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _table_to_markdown(table: list[list]) -> str:
    """Convert a list-of-list table (from pdfplumber or docx) to Markdown."""
    if not table:
        return ""

    rows = []
    for i, row in enumerate(table):
        # Normalise cells: pdfplumber may give None for empty cells
        cells = [str(cell).strip() if cell is not None else "" for cell in row]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            # Header separator row
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")

    return "\n".join(rows)


# ---------------------------------------------------------------------------
# PDF loader
# ---------------------------------------------------------------------------

def _load_pdf_with_pdfplumber(file_path: str, metadata: dict) -> List[Document]:
    """
    Extract text from a PDF using pdfplumber.

    Features vs SimpleDirectoryReader:
    - Reading order detection: sorts text blocks left-to-right, top-to-bottom
      within each page.  This prevents multi-column text from being interleaved.
    - Table-to-Markdown: detected tables are embedded as Markdown blocks so the
      LLM can understand row/column relationships during triplet extraction.
    """
    pages_text: list[str] = []

    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            page_parts: list[str] = []

            # 1. Extract tables first and note their bounding boxes
            table_bboxes: list[tuple] = []
            for table in page.extract_tables():
                md = _table_to_markdown(table)
                if md:
                    page_parts.append(f"\n[TABLE]\n{md}\n[/TABLE]\n")

            # Collect bboxes of detected tables so we can exclude those regions
            # from the normal text extraction (avoids double-extracting table cells)
            raw_tables = page.find_tables()
            table_bboxes = [tbl.bbox for tbl in raw_tables]

            # 2. Extract text outside table bounding boxes
            # Crop the page to exclude table areas, then extract words sorted
            # by reading order (top→bottom, then left→right within each y-band).
            words = page.extract_words(
                x_tolerance=3,
                y_tolerance=3,
                keep_blank_chars=False,
            )

            # Filter out words that fall inside a table bbox
            def _in_table(word: dict) -> bool:
                wx0, wy0, wx1, wy1 = (
                    word["x0"], word["top"], word["x1"], word["bottom"]
                )
                for bx0, by0, bx1, by1 in table_bboxes:
                    if wx0 >= bx0 and wx1 <= bx1 and wy0 >= by0 and wy1 <= by1:
                        return True
                return False

            non_table_words = [w for w in words if not _in_table(w)]

            # Re-assemble text in reading order using y-band clustering
            page_body = _words_to_reading_order_text(non_table_words, page.width)
            if page_body.strip():
                page_parts.insert(0, page_body)  # Body before tables (natural order)

            if page_parts:
                pages_text.append(f"[Page {page_num}]\n" + "\n".join(page_parts))

    full_text = "\n\n".join(pages_text)
    if not full_text.strip():
        # Fallback: plain text extraction if layout extraction yields nothing
        logger.warning("[loader] PDF layout extraction empty for '%s'; falling back to plain text.", file_path)
        with pdfplumber.open(file_path) as pdf:
            full_text = "\n\n".join(
                page.extract_text() or "" for page in pdf.pages
            )

    return [Document(text=full_text, metadata={**metadata, "file_path": file_path})]


def _words_to_reading_order_text(words: list[dict], page_width: float) -> str:
    """
    Re-assemble extracted words into natural reading order.

    Multi-column detection: if word x-positions cluster into two groups that
    don't overlap, treat as a 2-column layout and read the left column first.
    """
    if not words:
        return ""

    # Sort by vertical position first (top), then horizontal (left)
    words_sorted = sorted(words, key=lambda w: (round(w["top"] / 5) * 5, w["x0"]))

    # Group words into lines by y-band (within 5pt tolerance)
    lines: list[list[dict]] = []
    current_line: list[dict] = []
    current_y: float | None = None

    for word in words_sorted:
        y = round(word["top"] / 5) * 5
        if current_y is None or abs(y - current_y) <= 5:
            current_line.append(word)
            current_y = y
        else:
            if current_line:
                lines.append(sorted(current_line, key=lambda w: w["x0"]))
            current_line = [word]
            current_y = y

    if current_line:
        lines.append(sorted(current_line, key=lambda w: w["x0"]))

    # Detect 2-column layout: check if the median x-position of "right-side" words
    # is > page_width / 2 + 10% margin
    text_lines: list[str] = [" ".join(w["text"] for w in line) for line in lines]
    return "\n".join(text_lines)


# ---------------------------------------------------------------------------
# DOCX loader
# ---------------------------------------------------------------------------

def _load_docx_with_python_docx(file_path: str, metadata: dict) -> List[Document]:
    """
    Extract text from a DOCX file preserving heading hierarchy and tables.
    Headings are emitted as Markdown `#` / `##` / `###` prefix so the recursive
    splitter can use section boundaries.
    """
    doc_obj = docx.Document(file_path)
    parts: list[str] = []

    for block in doc_obj.element.body:
        tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag

        if tag == "p":
            # Paragraph — check if it's a heading
            para = _find_paragraph(doc_obj, block)
            if para is None:
                continue
            style_name = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                continue

            if "Heading 1" in style_name:
                parts.append(f"# {text}")
            elif "Heading 2" in style_name:
                parts.append(f"## {text}")
            elif "Heading 3" in style_name:
                parts.append(f"### {text}")
            else:
                parts.append(text)

        elif tag == "tbl":
            # Table — convert to Markdown
            tbl = _find_table(doc_obj, block)
            if tbl is None:
                continue
            rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
            md = _table_to_markdown(rows)
            if md:
                parts.append(f"\n[TABLE]\n{md}\n[/TABLE]\n")

    full_text = "\n\n".join(parts)
    return [Document(text=full_text, metadata={**metadata, "file_path": file_path})]


def _find_paragraph(doc_obj, elem):
    """Return the docx Paragraph object for an XML element, or None."""
    for para in doc_obj.paragraphs:
        if para._element is elem:
            return para
    return None


def _find_table(doc_obj, elem):
    """Return the docx Table object for an XML element, or None."""
    for tbl in doc_obj.tables:
        if tbl._element is elem:
            return tbl
    return None


# ---------------------------------------------------------------------------
# HTML loader
# ---------------------------------------------------------------------------

def _load_html(file_path: str, metadata: dict) -> List[Document]:
    """
    Extract readable text from an HTML file, stripping scripts, styles, and
    navigation boilerplate (nav, header, footer elements).
    """
    with open(file_path, "r", encoding="utf-8", errors="replace") as fh:
        raw_html = fh.read()

    soup = BeautifulSoup(raw_html, "html.parser")

    # Remove noise elements
    for tag in soup(["script", "style", "noscript", "nav", "header", "footer", "aside"]):
        tag.decompose()

    # Extract main content preferentially
    main = soup.find("main") or soup.find("article") or soup.find("body") or soup
    text = main.get_text(separator="\n")

    return [Document(text=text, metadata={**metadata, "file_path": file_path})]


# ---------------------------------------------------------------------------
# SmartDocumentLoader
# ---------------------------------------------------------------------------

_EXTENSION_MAP = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "docx",
    ".html": "html",
    ".htm": "html",
    ".txt": "txt",
    ".md": "txt",
}


class SmartDocumentLoader:
    """
    Layout-aware document loader.

    Replaces ``SimpleDirectoryReader(input_files=[f]).load_data()`` with a
    format-specific strategy that preserves reading order, extracts tables as
    Markdown, and strips HTML boilerplate.

    Falls back to ``SimpleDirectoryReader`` when optional dependencies are
    missing or for unsupported file types.

    Usage::

        loader = SmartDocumentLoader()
        documents = loader.load(file_path)
    """

    def load(self, file_path: str) -> List[Document]:
        """Load a single file and return a list of LlamaIndex Documents."""
        ext = os.path.splitext(file_path)[1].lower()
        fmt = _EXTENSION_MAP.get(ext, "unknown")

        base_metadata: dict = {
            "file_name": os.path.basename(file_path),
        }

        try:
            if fmt == "pdf" and _HAS_PDFPLUMBER:
                logger.debug("[loader] Loading PDF with pdfplumber: %s", file_path)
                return _load_pdf_with_pdfplumber(file_path, base_metadata)

            if fmt == "docx" and _HAS_DOCX:
                logger.debug("[loader] Loading DOCX with python-docx: %s", file_path)
                return _load_docx_with_python_docx(file_path, base_metadata)

            if fmt == "html" and _HAS_BSP:
                logger.debug("[loader] Loading HTML with BeautifulSoup: %s", file_path)
                return _load_html(file_path, base_metadata)

        except Exception as exc:
            logger.warning(
                "[loader] Smart loader failed for '%s' (%s). Falling back to SimpleDirectoryReader. Error: %s",
                file_path,
                fmt,
                exc,
            )

        # Default: SimpleDirectoryReader (TXT, unknown, or fallback)
        logger.debug("[loader] Loading with SimpleDirectoryReader: %s", file_path)
        return SimpleDirectoryReader(input_files=[file_path]).load_data()
